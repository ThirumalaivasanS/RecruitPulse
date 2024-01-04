import requests
from bs4 import BeautifulSoup
import docx
import os
import zipfile
import io
import json
from concurrent.futures import ThreadPoolExecutor
import re

def scrape_resumes(job_title, state):
    url = f"https://www.postjobfree.com/resumes?q=&n=&t={job_title.replace(' ', '+')}&d=&l={state.replace(' ', '+')}&radius=25&r=100"
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    title_tags = soup.find_all('h3', attrs={'class': 'itemTitle'})
    links = ["https://www.postjobfree.com" + title_tag.a['href'] for title_tag in title_tags]
    return links

def scrape_resume_content(link):
    response = requests.get(link)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        resume_content = soup.find('div', class_='normalText')  
        if resume_content:
            return resume_content.text.strip()
        else:
            print(f"Unable to extract from: {link}")
    else:
        print(f"Failed to fetch: {link}")
    return None

def load_conditions_from_json(file_path):
    with open(file_path, 'r') as json_file:
        data = json.load(json_file)
    return data

def save_resume_as_docx(content, job_role, output_folder, must_have, good_to_have, file_suffix):
    file_name = f"{job_role}_resume_{file_suffix}.docx"
    file_path = os.path.join(output_folder, file_name)
    doc = docx.Document()
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(content)
    
    words = re.findall(r'\b\w+\b|\bC\+\+\b', content)
    
    for word in words:
        if word.lower() in must_have:
            start_index = content.lower().find(word.lower())
            end_index = start_index + len(word)
            run.add_text(content[start_index:end_index])  # Highlighting the must-have word
            font = run.font
            font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            font.bold = True  
        
        elif word.lower() in good_to_have:
            start_index = content.lower().find(word.lower())
            end_index = start_index + len(word)
            run.add_text(content[start_index:end_index])  # Highlighting the good-to-have word
            font = run.font
            font.highlight_color = docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN
    
    doc.save(file_path)
    return file_path



def process_links(links):
    with ThreadPoolExecutor(max_workers=10) as executor:
        results = list(executor.map(scrape_resume_content, links))
    return results

if __name__ == '__main__':
    output_folder = r'C:\Users\Admin\Videos\output'  
    json_file_path = r'C:\Users\Admin\Videos\demo_j.json'  
    conditions = load_conditions_from_json(json_file_path)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        for role in conditions:
            job_role = role['job_title']
            must_have_skills_list = role['must_have']
            good_have_skills_list = role['good_to_have']
            
            resumes_data = scrape_resumes(job_role, role['locations'][0]) 
            resume_contents = process_links(resumes_data)

            for idx, resume_content in enumerate(resume_contents):
                file_path = save_resume_as_docx(resume_content, job_role, output_folder, must_have_skills_list, good_have_skills_list, idx)
                if file_path:
                    zip_file.write(file_path, os.path.basename(file_path))

    zip_filename = os.path.join(output_folder, 'resumes.zip')
    with open(zip_filename, 'wb') as f:
        f.write(zip_buffer.getvalue())

    print(f"All resumes downloaded and saved in {zip_filename}")