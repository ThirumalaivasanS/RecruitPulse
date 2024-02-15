from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from models.keyword import extract_keywords_and_save
from models.ranking import rank_resumes, process_resume_link
from models.scraping import scrape_and_get_links
import os
from docx import Document
import pandas as pd
import json
import time

app = Flask(__name__)

def save_to_doc(content, file_path):
    document = Document()
    document.add_paragraph(content)
    document.save(file_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_file():
    global keywords_result
    
    if 'file' not in request.files:
        return render_template('index.html', error='No file part')

    file = request.files['file']

    if file.filename == '':
        return render_template('index.html', error='No selected file')

    if file:
        filename = secure_filename(file.filename)
        file_path = 'uploads/' + filename
        file.save(file_path)

        keywords_result = extract_keywords_and_save(file_path)
        links_result = scrape_and_get_links([keywords_result])
        ranking_result = rank_resumes(file_path, links_result)
        ranking_result = ranking_result[ranking_result['Similarity'] >= 0.6] 
        ranking_result = ranking_result.sort_values(by='Similarity', ascending=False)
        ranking_result = ranking_result.round(2)

        time.sleep(3)
        response_data = {'status': 'done'}


        return render_template('index.html', ranking_result=ranking_result.to_dict('records'))

    return render_template('index.html', error='Error processing file')

@app.route('/download/<path:link>')
def download(link):
    name, content = process_resume_link(link)
    doc = Document()
    doc.add_paragraph(content)
    
    safe_filename = secure_filename(link)
    
    download_path = os.path.join('downloaded_resumes', f'{name}_resume.doc')
    doc.save(download_path)
    return send_file(download_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
